"""Build a Google Docs-compatible DOCX with experiment images embedded."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent
REPORTS = ROOT / "reports"
OUTPUT = REPORTS / "Prompt_and_Model_Experimentation_Report.docx"


def add_picture(document: Document, relative_path: str, caption: str) -> None:
    document.add_picture(str(REPORTS / relative_path), width=Inches(6.5))
    paragraph = document.paragraphs[-1]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph = document.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.style = document.styles["Caption"]


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, headers, strict=True):
        cell.text = value
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row, strict=True):
            cell.text = value


def main() -> None:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    document.add_heading("Prompt and Model Experimentation Summary", level=0)

    document.add_heading("1. Prompt experimentation", level=1)
    document.add_paragraph(
        "The initial extraction prompt did not explicitly require descriptions "
        "to be copied from the resume. Consequently, the model sometimes "
        "paraphrased experience and project descriptions."
    )
    document.add_paragraph("The following rule was added to the improved prompt:")
    quote = document.add_paragraph()
    quote.style = document.styles["Quote"]
    quote.add_run(
        "Copy experience and project descriptions word-for-word as written. "
        "Do not summarise, rewrite, improve, or paraphrase them."
    )
    document.add_paragraph(
        "Both prompts were tested using the same five technology resumes and "
        "the same Gemma 4 31B model."
    )
    add_table(
        document,
        ["Prompt", "Exact description copy rate"],
        [
            ["Scenario A: original prompt", "20%"],
            ["Scenario B: improved verbatim prompt", "45%"],
        ],
    )
    document.add_paragraph(
        "The exact copy rate increased by 25 percentage points after adding the "
        "explicit instruction. This shows that prompt wording materially "
        "affected extraction behaviour."
    )
    add_picture(
        document,
        "gemma4_prompt_ab/description_quality_metrics.png",
        "Figure 1. Exact description copy rate by prompt version.",
    )
    add_table(
        document,
        ["Prompt", "Average latency"],
        [
            ["Scenario A: original prompt", "76.0 seconds"],
            ["Scenario B: improved verbatim prompt", "96.0 seconds"],
        ],
    )
    add_picture(
        document,
        "gemma4_prompt_ab/average_latency.png",
        "Figure 2. Average latency by prompt version.",
    )

    document.add_heading("2. Model comparison using the improved prompt", level=1)
    document.add_paragraph(
        "The improved verbatim prompt was kept fixed while the model was changed "
        "from Gemma 4 31B to Gemini 3.5 Flash. Both models processed the same "
        "five technology resumes."
    )
    add_table(
        document,
        ["Model", "Successful runs", "Average latency", "Exact copy rate*"],
        [
            ["Gemma 4 31B", "5/5", "96.0 seconds", "33.3%"],
            ["Gemini 3.5 Flash", "5/5", "14.3 seconds", "91.7%"],
        ],
    )
    document.add_paragraph(
        "*For the direct model comparison, exact copy rate is calculated over "
        "resumes where the model returned at least one description."
    )
    document.add_paragraph(
        "Gemini 3.5 Flash was approximately 6.7 times faster and produced "
        "substantially more descriptions that matched the resume word-for-word. "
        "It was therefore selected for the implemented description-extraction "
        "configuration."
    )
    add_picture(
        document,
        "model_comparison/average_latency.png",
        "Figure 3. Average latency comparison using the improved prompt.",
    )
    add_picture(
        document,
        "model_comparison/exact_copy_rate.png",
        "Figure 4. Exact description copy rate comparison using the improved prompt.",
    )

    document.add_heading("3. Example implementation output", level=1)
    document.add_paragraph(
        "A complete extracted JSON example is included separately as "
        "gemini35_verbatim/sample_extracted_resume.json. Direct identifiers and "
        "identifying dates were redacted; the extracted structure and "
        "non-identifying fields are unchanged."
    )

    document.add_heading("4. Reproducibility", level=1)
    command = document.add_paragraph()
    command.style = document.styles["No Spacing"]
    command.add_run(
        ".venv/bin/python -u experiments/run_description_ab.py "
        "--model gemini-3.5-flash --versions v002 "
        "--experiment gemini35_verbatim_raw "
        "--results gemini35_verbatim.jsonl\n"
        ".venv/bin/python experiments/analyze_model_comparison.py"
    ).font.name = "Courier New"

    document.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
