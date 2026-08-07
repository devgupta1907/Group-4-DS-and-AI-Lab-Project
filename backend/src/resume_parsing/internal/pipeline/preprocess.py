"""Stage 2 — turn an accepted document into per-page work units.

Vision path: render each page to a PNG. Deliberately minimal — the model reads
the page directly, so there is no OCR stage and no cropping. Light enhancement
runs only when a page actually looks poor, never as a blanket step.

Nothing here writes to disk. Page images exist as bytes inside the request and
are released when the parse finishes, which is what makes the "temporary
artifacts are deleted immediately" guarantee hold without a cleanup job.
"""

from __future__ import annotations

import io

import fitz  # PyMuPDF
from PIL import Image, ImageOps

from src.core.config import get_settings
from src.resume_parsing.errors import UnreadableDocument
from src.resume_parsing.internal.domain import PageArtifact, SourceDocument
from src.resume_parsing.schemas import ParseRoute

# Below this mean luminance spread, a page is flat enough that contrast
# normalisation is worth applying.
_LOW_CONTRAST_STDDEV = 45.0


def to_pages(document: SourceDocument) -> list[PageArtifact]:
    if document.route is ParseRoute.TEXT:
        return _text_pages(document)
    return _image_pages(document)


def text_artifact(text: str) -> list[PageArtifact]:
    """Wrap externally extracted document text for the ordinary LLM stage."""
    cleaned = _clean(text)
    if not cleaned:
        raise UnreadableDocument("Document preprocessing returned no readable text.")
    return [PageArtifact(index=0, text=cleaned)]


# ------------------------------------------------------------------- text path --


def _text_pages(document: SourceDocument) -> list[PageArtifact]:
    if document.media_type.endswith("wordprocessingml.document"):
        return [PageArtifact(index=0, text=_docx_text(document.content))]

    try:
        pdf = fitz.open(stream=document.content, filetype="pdf")
    except Exception as exc:
        raise UnreadableDocument() from exc
    try:
        return [
            PageArtifact(index=i, text=_clean(pdf[i].get_text("text")))
            for i in range(pdf.page_count)
        ]
    finally:
        pdf.close()


def _docx_text(content: bytes) -> str:
    from docx import Document  # imported lazily: only the DOCX path needs it

    try:
        docx = Document(io.BytesIO(content))
    except Exception as exc:
        raise UnreadableDocument("That DOCX file could not be opened.") from exc

    blocks = [p.text for p in docx.paragraphs]
    for table in docx.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text for cell in row.cells))
    return _clean("\n".join(blocks))


def _clean(text: str) -> str:
    lines = (line.strip() for line in text.replace("\r\n", "\n").split("\n"))
    return "\n".join(line for line in lines if line)


# ----------------------------------------------------------------- vision path --


def _image_pages(document: SourceDocument) -> list[PageArtifact]:
    if document.media_type.startswith("image/"):
        return [PageArtifact(index=0, image_png=_normalise_image(document.content))]

    settings = get_settings()
    try:
        pdf = fitz.open(stream=document.content, filetype="pdf")
    except Exception as exc:
        raise UnreadableDocument() from exc
    try:
        pages = []
        for i in range(pdf.page_count):
            pixmap = pdf[i].get_pixmap(dpi=settings.resume_render_dpi)
            pages.append(
                PageArtifact(index=i, image_png=_enhance_if_poor(pixmap.tobytes("png")))
            )
        return pages
    finally:
        pdf.close()


def _normalise_image(content: bytes) -> bytes:
    """Re-encode an uploaded image as PNG so the provider sees one format."""
    try:
        image = Image.open(io.BytesIO(content))
    except Exception as exc:
        raise UnreadableDocument("That image could not be opened.") from exc
    return _enhance_if_poor(_to_png(image.convert("RGB")))


def _enhance_if_poor(png: bytes) -> bytes:
    """Autocontrast a page only when its luminance is genuinely flat."""
    try:
        image = Image.open(io.BytesIO(png))
        grey = image.convert("L")
        if _stddev(grey) >= _LOW_CONTRAST_STDDEV:
            return png
        return _to_png(ImageOps.autocontrast(image.convert("RGB"), cutoff=1))
    except Exception:
        # Enhancement is opportunistic; a failure here must not fail the parse.
        return png


def _stddev(grey: Image.Image) -> float:
    histogram = grey.histogram()
    total = sum(histogram) or 1
    mean = sum(i * n for i, n in enumerate(histogram)) / total
    variance = sum(n * (i - mean) ** 2 for i, n in enumerate(histogram)) / total
    return variance**0.5


def _to_png(image: Image.Image) -> bytes:
    buffer = io.BytesIO()
    image.save(buffer, format="PNG", optimize=True)
    return buffer.getvalue()
